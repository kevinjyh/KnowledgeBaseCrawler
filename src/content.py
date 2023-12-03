# content.py
import pdfplumber
from docx import Document
import openpyxl
import os

class ContentExtractor:
    """ 負責根據檔案類型提取檔案內容的類別 """

    @staticmethod
    def get_local_file_content(self, file_path):
        full_path = os.path.abspath(file_path)
        file_extension = os.path.splitext(file_path)[1].lower().strip('.')
        content = {'text': [], 'tables': []}
        if file_extension in ['pdf']:
            content['text'] = self.extract_text_from_pdf(full_path)
            content['tables'] = self.extract_tables_from_pdf(full_path)
        elif file_extension in ['docx', 'doc']:
            content['text'] = self.extract_text_from_docx(full_path)
            content['tables'] = self.extract_tables_from_docx(full_path)
        elif file_extension in ['xlsx', 'xls']:
            content = self.extract_text_from_xlsx(full_path)
        elif file_extension in ['txt', 'md', 'html']:
            content = self.extract_text_from_txt(full_path)
        else:
            content = None  # f"****尚未支援{file_extension}類型的檔案****"
        
        # 在此生成或獲取緩存中的 file_id
        with self.lock:
            file_id = self.generate_file_id(file_path)

        return {
            'file_path': full_path,  # 使用完整路徑
            'file_id': file_id,       # 添加 file_id
            'content': content
        }

    @staticmethod
    def extract_text_from_pdf(file_path):
        try:
            with pdfplumber.open(file_path) as pdf:
                return [{'text': page.extract_text(), 'page_number': i+1} for i, page in enumerate(pdf.pages) if page.extract_text()]
        except Exception as e:
            print(f"Error in processing PDF {file_path}: {e}")
            return []

    @staticmethod
    def extract_tables_from_pdf(file_path):
        with pdfplumber.open(file_path) as pdf:
            tables = []
            for page in pdf.pages:
                page_tables = page.extract_tables()
                for table in page_tables:
                    tables.append([row for row in table])
            return tables

    @staticmethod
    def extract_text_from_docx(file_path):
        doc = Document(file_path)
        return [{'text': paragraph.text, 'page_number': i+1} for i, paragraph in enumerate(doc.paragraphs) if paragraph.text]

    @staticmethod
    def extract_tables_from_docx(file_path):
        doc = Document(file_path)
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            tables.append(table_data)
        return tables

    @staticmethod
    def extract_text_from_xlsx(file_path):
        workbook = openpyxl.load_workbook(file_path, read_only=True)
        data = []
        for sheet in workbook.worksheets:
            sheet_data = {
                'sheet_name': sheet.title,
                'contents': '\n'.join(
                    (str(cell.value) if cell.value is not None else '') for row in sheet for cell in row
                )
            }
            data.append(sheet_data)
        return data

    @staticmethod
    def extract_text_from_txt(file_path):
        with open(file_path, 'r', encoding='utf-8') as file:
            return {'text': file.read()}