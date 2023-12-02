import concurrent.futures
import hashlib
import os
import logging
from datetime import datetime
import threading
from functools import lru_cache
import json
import pdfplumber
from docx import Document
import openpyxl
import fnmatch

class FileExtractor:
    def __init__(self, config_path='src/config.json'):
        self.config = self.load_config(config_path)
        self.local_path = self.config["local_path"]
        self.ignore_patterns = self.config["ignore"]
        self.max_size_mb = self.config["max_size_mb"]
        self.base_output_file_name = os.path.splitext(self.config["output_file_name"])[0]  # 不包含副檔名
        # 初始化緩存大小和鎖
        self.file_id_cache = lru_cache(maxsize=10000)
        self.lock = threading.Lock()
        # 設置日誌
        logging.basicConfig(filename='src/updated.log',  # 根目錄下的日誌文件
                            level=logging.INFO,
                            format='%(asctime)s - %(levelname)s - %(message)s')

    def load_existing_data(self):
        existing_data = []
        # 获取当前脚本文件所在的目录
        current_dir = os.path.dirname(os.path.abspath(__file__))
        
        # 在当前目录下搜索 JSON 文件
        for filename in os.listdir(current_dir):
            if filename.startswith(self.base_output_file_name + '_') and filename.endswith(".json"):
                # 使用当前目录和文件名构造完整的文件路径
                file_path = os.path.join(current_dir, filename)
                with open(file_path, 'r', encoding='utf-8') as file:
                    data = json.load(file)
                    existing_data.extend(data)
        return existing_data

    def file_id_exists(self, file_id, existing_data):
        return any(item['file_id'] == file_id for item in existing_data)

    def process_files(self):
        # 获取当前时间，格式化为完整日期和24小时制时间
        current_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        # 创建分隔线，包含当前时间
        separator = '=' * 20 + f' {current_time} ' + '=' * 20
        
        # 在日志文件中记录分隔线
        logging.info(separator)

        existing_data = self.load_existing_data()
        crawled_data = []
        files = self.get_files_in_directory(self.local_path)

        updated_files = []
        files = self.get_files_in_directory(self.local_path)

        with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
            futures = {executor.submit(self.get_local_file_content, file): file for file in files}
            for future in concurrent.futures.as_completed(futures):
                file = futures[future]
                try:
                    data = future.result()
                    if not self.file_id_exists(data['file_id'], existing_data):
                        print(f"Crawling {file}")
                        crawled_data.append(data)
                        updated_files.append(file)
                        logging.info(f'Updated file: {file}')
                    else:
                        print('.', end='')
                except Exception as exc:
                    logging.error(f'File processing generated an exception: {file}, {exc}')

        self.write_to_file(crawled_data, self.base_output_file_name, self.max_size_mb)

    def get_local_file_content(self, file_path):
        full_path = os.path.abspath(file_path)
        file_extension = os.path.splitext(file_path)[1].lower().strip('.')
        content = {'text': [], 'tables': []}
        if file_extension in ['pdf']:
            content['text'] = self.extract_text_from_pdf(full_path)
            content['tables'] = self.extract_tables_from_pdf(full_path)
        elif file_extension in ['docx', 'doc']:
            content['text'] = self.extract_text_from_docx(full_path)
            content['tables'] = self.extract_tables_from_docx(full_path)
        elif file_extension in ['xlsx', 'xls']:
            content = self.extract_text_from_xlsx(full_path)
        elif file_extension in ['txt', 'md', 'html']:
            content = self.extract_text_from_txt(full_path)
        else:
            content = None  # f"****尚未支援{file_extension}類型的檔案****"
        
        # 在此生成或獲取緩存中的 file_id
        with self.lock:
            file_id = self.generate_file_id(file_path)

        return {
            'file_path': full_path,  # 使用完整路徑
            'file_id': file_id,       # 添加 file_id
            'content': content
        }

    @staticmethod
    @lru_cache(maxsize=10000)
    def generate_file_id(file_path):
        try:
            # 獲取文件的絕對路徑、創建日期和修改日期
            absolute_path = os.path.abspath(file_path)
            creation_time = os.path.getctime(file_path)
            modification_time = os.path.getmtime(file_path)
            # 串聯這些信息
            data = f"{absolute_path}{creation_time}{modification_time}"
            # 使用 SHA-1 生成雜湊
            return hashlib.sha1(data.encode()).hexdigest()
        except Exception as e:
            print(f"Error generating file ID for {file_path}: {e}")
            return None

    @staticmethod
    def load_config(config_path):
        with open(config_path, 'r', encoding='utf-8') as config_file:
            return json.load(config_file)

    @staticmethod
    def extract_text_from_pdf(file_path):
        try:
            with pdfplumber.open(file_path) as pdf:
                return [{'text': page.extract_text(), 'page_number': i+1} for i, page in enumerate(pdf.pages) if page.extract_text()]
        except Exception as e:
            print(f"Error in processing PDF {file_path}: {e}")
            return []

    @staticmethod
    def extract_tables_from_pdf(file_path):
        with pdfplumber.open(file_path) as pdf:
            tables = []
            for page in pdf.pages:
                page_tables = page.extract_tables()
                for table in page_tables:
                    tables.append([row for row in table])
            return tables

    @staticmethod
    def extract_text_from_docx(file_path):
        doc = Document(file_path)
        return [{'text': paragraph.text, 'page_number': i+1} for i, paragraph in enumerate(doc.paragraphs) if paragraph.text]

    @staticmethod
    def extract_tables_from_docx(file_path):
        doc = Document(file_path)
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            tables.append(table_data)
        return tables

    @staticmethod
    def extract_text_from_xlsx(file_path):
        workbook = openpyxl.load_workbook(file_path, read_only=True)
        data = []
        for sheet in workbook.worksheets:
            sheet_data = {
                'sheet_name': sheet.title,
                'contents': '\n'.join(
                    (str(cell.value) if cell.value is not None else '') for row in sheet for cell in row
                )
            }
            data.append(sheet_data)
        return data

    @staticmethod
    def extract_text_from_txt(file_path):
        with open(file_path, 'r', encoding='utf-8') as file:
            return {'text': file.read()}

    def get_files_in_directory(self, directory_path):
        files = []
        for root, dirs, files_in_dir in os.walk(directory_path):
            for filename in files_in_dir:
                file_path = os.path.join(root, filename)
                if os.path.isfile(file_path) and not any(fnmatch.fnmatch(filename, pattern) for pattern in self.ignore_patterns):
                    files.append(file_path)
        return files

    def write_to_file(self, data, base_filename, max_size_mb, encoding='utf-8'):
        # 获取当前运行的 Python 文件的目录
        directory = os.path.dirname(os.path.abspath(__file__))

        # 初始文件索引
        file_index = 1
        # 数据分块大小
        chunk_size = 1024 * 1024 * max_size_mb  # 将MB转换为字节
        # 初始化当前块
        current_chunk = []
        current_size = 0

        for item in data:
            item_size = len(json.dumps(item, ensure_ascii=False).encode(encoding))
            if current_size + item_size > chunk_size:
                # 当前块的大小已经超过限制，写入文件
                filename = f'{base_filename}_{file_index}.json'
                full_path = os.path.join(directory, filename)  # 构建完整的文件路径
                with open(full_path, 'w', encoding=encoding) as f:
                    json.dump(current_chunk, f, ensure_ascii=False, indent=4)
                # 重置当前块和大小计数器
                current_chunk = [item]
                current_size = item_size
                file_index += 1
            else:
                # 将当前项目添加到块中
                current_chunk.append(item)
                current_size += item_size

        # 写入最后一个文件（如果有剩余数据）
        if current_chunk:
            filename = f'{base_filename}_{file_index}.json'
            full_path = os.path.join(directory, filename)  # 构建完整的文件路径
            with open(full_path, 'w', encoding=encoding) as f:
                json.dump(current_chunk, f, ensure_ascii=False, indent=4)

    def crawl(self):
        # 使用 process_files 來處理所有文件
        self.process_files()    # 為了進行Unit test 所以保留這行


# Main execution
if __name__ == "__main__":
    FileExtractor().crawl()
    